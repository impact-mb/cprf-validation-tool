import io
import re
import zipfile
from datetime import datetime
from zoneinfo import ZoneInfo
from pathlib import Path
import hashlib
import hmac
import uuid

import gspread
from docx import Document

import pandas as pd
import streamlit as st


APP_VERSION = "2026.01.02"
APP_OWNER = "Magic Bus Impact Team"


# --------------- Authentication Helpers --------------- #
HASH_SCHEME = "pbkdf2_sha256"


def get_login_users():
    """Read all configured users from Streamlit Secrets.

    Expected secrets structure:

        [users.Narendra]
        password_hash = "pbkdf2_sha256$..."
        role = "admin"

        [users.MB_FPD]
        password_hash = "pbkdf2_sha256$..."
        role = "user"

    Plaintext passwords are never stored.
    """
    try:
        users = st.secrets["users"]
        result = {}
        for username, config in users.items():
            password_hash = str(config.get("password_hash", "")).strip()
            role = str(config.get("role", "user")).strip().lower()
            if password_hash:
                result[str(username)] = {
                    "password_hash": password_hash,
                    "role": role or "user",
                }
        return result
    except Exception:
        return {}


def verify_password(password: str, stored_hash: str) -> bool:
    """Verify a password against a salted PBKDF2-SHA256 hash.

    Expected format:
        pbkdf2_sha256$ITERATIONS$SALT_HEX$DIGEST_HEX
    """
    try:
        scheme, iterations_text, salt_hex, digest_hex = stored_hash.split("$", 3)
        if scheme != HASH_SCHEME:
            return False

        iterations = int(iterations_text)
        if iterations < 100_000:
            return False

        salt = bytes.fromhex(salt_hex)
        expected_digest = bytes.fromhex(digest_hex)
        actual_digest = hashlib.pbkdf2_hmac(
            "sha256",
            password.encode("utf-8"),
            salt,
            iterations,
        )
        return hmac.compare_digest(actual_digest, expected_digest)
    except (ValueError, TypeError, AttributeError):
        return False


def authenticate(username: str, password: str):
    """Authenticate a configured user and return account metadata on success."""
    users = get_login_users()
    if not users:
        return None

    # Match usernames exactly while avoiding accidental whitespace issues.
    account = users.get(username)
    if account is None:
        return None

    if verify_password(password, account["password_hash"]):
        return {"username": username, "role": account.get("role", "user")}
    return None


def render_login_page():
    """Render the branded landing/login layer."""
    st.markdown(
        f"""
        <style>
        .stApp {
            background: linear-gradient(135deg, #f7f8fa 0%, #eef2f6 100%);
        }
        .login-hero {
            max-width: 760px;
            margin: 6vh auto 1.2rem auto;
            text-align: center;
        }
        .login-title {
            font-size: 2.35rem;
            font-weight: 750;
            margin-bottom: .25rem;
        }
        .login-subtitle {
            font-size: 1.05rem;
            color: #4b5563;
            margin-bottom: .35rem;
        }
        .release-pill {
            display: inline-block;
            padding: .3rem .7rem;
            border: 1px solid #d1d5db;
            border-radius: 999px;
            font-size: .85rem;
            color: #374151;
            background: white;
        }
        .login-foot {
            text-align:center;
            color:#6b7280;
            font-size:.82rem;
            margin-top:1.2rem;
        }
        </style>
        <div class="login-hero">
            <div class="login-title">CPRF Validation Tool</div>
            <div class="login-subtitle">Data Quality & Validation Workspace</div>
            <span class="release-pill">Release v{APP_VERSION} · {APP_OWNER}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    left, centre, right = st.columns([1.2, 1, 1.2])
    with centre:
        with st.form("login_form", clear_on_submit=False):
            st.markdown("### Sign in")
            st.caption("Authorized Magic Bus users only")
            username = st.text_input("Username", placeholder="Enter username")
            password = st.text_input(
                "Password", type="password", placeholder="Enter password"
            )
            submitted = st.form_submit_button(
                "Login", use_container_width=True, type="primary"
            )

        if submitted:
            configured_users = get_login_users()
            if not configured_users:
                st.error(
                    "Login is not configured. Add the [users.<username>] sections "
                    "with password_hash and role to Streamlit Secrets."
                )
            else:
                account = authenticate(username.strip(), password)
                if account:
                    st.session_state["authenticated"] = True
                    st.session_state["authenticated_user"] = account["username"]
                    st.session_state["authenticated_role"] = account["role"]
                    log_successful_login(account["username"], account["role"])
                    st.rerun()
                else:
                    st.error("Invalid username or password.")

    st.markdown(
        '<div class="login-foot">Internal utility · Magic Bus Impact Team</div>',
        unsafe_allow_html=True,
    )


def render_authenticated_header():
    """Top bar shown after successful login."""
    title_col, user_col, logout_col = st.columns([6, 2, 1])
    with title_col:
        st.markdown(
            f"### CPRF Validation Tool <span style='font-size:.78rem; "
            f"font-weight:500; color:#6b7280;'>v{APP_VERSION}</span>",
            unsafe_allow_html=True,
        )
        st.caption(f"Released by {APP_OWNER}")
    with user_col:
        user = st.session_state.get("authenticated_user", "User")
        role = st.session_state.get("authenticated_role", "user")
        st.caption("Signed in as")
        st.write(f"**{user}** · {role.title()}")
    with logout_col:
        if st.button("Logout", use_container_width=True):
            st.session_state["authenticated"] = False
            st.session_state.pop("authenticated_user", None)
            st.session_state.pop("authenticated_role", None)
            st.rerun()

    st.divider()


# --------------- Google Sheet Login Audit --------------- #
def log_successful_login(username: str, role: str) -> bool:
    """Append one successful login event to the configured Google Sheet.

    Expected Streamlit Secrets:

        [login_sheet]
        spreadsheet_id = "..."
        worksheet_name = "Login_Log"

        [gcp_service_account]
        ...Google service account fields...

    Logging is deliberately fail-safe: authentication is never blocked when
    Google Sheets is temporarily unavailable.
    """
    try:
        spreadsheet_id = str(st.secrets["login_sheet"]["spreadsheet_id"]).strip()
        worksheet_name = str(
            st.secrets["login_sheet"].get("worksheet_name", "Login_Log")
        ).strip()
        credentials = dict(st.secrets["gcp_service_account"])

        client = gspread.service_account_from_dict(credentials)
        worksheet = client.open_by_key(spreadsheet_id).worksheet(worksheet_name)

        login_dt = datetime.now(ZoneInfo("Asia/Kolkata"))
        event_id = str(uuid.uuid4())
        worksheet.append_row(
            [
                event_id,
                username,
                role,
                login_dt.strftime("%d-%m-%Y"),
                login_dt.strftime("%H:%M:%S"),
                login_dt.strftime("%Y-%m-%d %H:%M:%S"),
                APP_VERSION,
            ],
            value_input_option="USER_ENTERED",
        )
        print(
            f"Google Sheet login audit recorded: user={username}, "
            f"role={role}, event_id={event_id}"
        )
        return True
    except Exception as exc:
        print(f"Google Sheet login logging failed: {exc}")
        return False


# --------------- Documentation Helpers --------------- #
DOCS_DIR = Path("docs")
MANUAL_PATH = DOCS_DIR / "CPRF_VALIDATION_TOOL_MANUAL.md"
RELEASE_NOTES_PATH = DOCS_DIR / "RELEASE_NOTES.md"


def read_text_file(path: Path) -> str:
    """Read a UTF-8 documentation file without breaking the app if absent."""
    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        return ""


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Create a simple, shareable Word copy from the maintained Markdown manual."""
    document = Document()
    document.core_properties.title = "CPRF Validation Tool - User & Technical Manual"
    document.core_properties.subject = f"CPRF Validation Tool v{APP_VERSION}"
    document.core_properties.author = APP_OWNER

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph()
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            document.add_paragraph(text, style="List Number")
        elif stripped.startswith("```"):
            continue
        elif stripped.startswith("|"):
            # Markdown tables remain readable as text in the generated Word copy.
            document.add_paragraph(stripped)
        else:
            clean = stripped.replace("**", "").replace("`", "")
            document.add_paragraph(clean)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def render_manual_page():
    st.subheader("User & Technical Manual")
    st.caption(
        "The Markdown file in the repository is the master copy. "
        "Update it whenever the tool changes."
    )
    manual = read_text_file(MANUAL_PATH)
    if not manual:
        st.warning(
            "Manual not found. Add docs/CPRF_VALIDATION_TOOL_MANUAL.md to the repository."
        )
        return

    st.markdown(manual)
    st.divider()
    col_md, col_docx = st.columns(2)
    file_date = datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%d%m%Y")
    with col_md:
        st.download_button(
            "Download Manual (.md)",
            data=manual.encode("utf-8"),
            file_name=f"CPRF_Validation_Tool_Manual_v{APP_VERSION}_{file_date}.md",
            mime="text/markdown",
            use_container_width=True,
        )
    with col_docx:
        st.download_button(
            "Download Manual (.docx)",
            data=markdown_to_docx_bytes(manual),
            file_name=f"CPRF_Validation_Tool_Manual_v{APP_VERSION}_{file_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )


def render_release_notes_page():
    st.subheader("Release Notes")
    notes = read_text_file(RELEASE_NOTES_PATH)
    if not notes:
        st.warning("Release notes not found. Add docs/RELEASE_NOTES.md to the repository.")
        return
    st.markdown(notes)


# --------------- Core Processing Function --------------- #
def process_excel(df: pd.DataFrame) -> pd.DataFrame:
    """
    Validate CPRF Excel data and return flagged DataFrame.

    Rules implemented:
    - Drop rows containing "Applied filters:" in any column
    - Drop columns DOCUMENTTYPE and DOCUMENTNO (if present)
    - School UDISE length check (must be exactly 11 characters)
    - DATE OF BIRTH format/missing check
    - AGE range check (only for PROGRAMSUBTYPENAME == 'ADOLOSCENT')
    - Parent Consent check
    - P_Age check (if ANY 0 exists → error)
    - RELIGIONNAME missing check
    - Total_Errors = sum of all error flags
    """

    df = df.copy()

    # ---------- DROP ROWS containing "Applied filters:" anywhere ---------- #
    df = df[
        ~df.apply(
            lambda row: row.astype(str)
            .str.contains("Applied filters:", case=False, na=False)
            .any(),
            axis=1,
        )
    ].reset_index(drop=True)

    # ---------- DROP UNUSED COLUMNS IF PRESENT ---------- #
    df = df.drop(columns=["DOCUMENTTYPE", "DOCUMENTNO"], errors="ignore")

    # ---------- CHECK REQUIRED COLUMNS ---------- #
    required_cols = [
        "School UDISE",
        "DATE OF BIRTH",
        "AGE",
        "PROGRAMSUBTYPENAME",
        "Parent Consent",
        "P_Age",
        "RELIGIONNAME",
        "ProgramLaunchName",
    ]
    missing_cols = [c for c in required_cols if c not in df.columns]
    if missing_cols:
        raise ValueError(f"Missing required columns: {', '.join(missing_cols)}")

    # ---------- INITIALIZE FLAG COLUMNS ---------- #
    df["ERROR_SCHOOL_UDISE"] = 0
    df["ERROR_DOB_FORMAT"] = 0
    df["ERROR_AGE_RANGE"] = 0
    df["ERROR_PARENT_CONSENT"] = 0
    df["ERROR_P_AGE"] = 0
    df["ERROR_RELIGIONNAME"] = 0

    # ---------- 1. School UDISE Length Check (len != 11) ---------- #
    udise_str = (
        df["School UDISE"]
        .str.strip()
    )
    df.loc[udise_str.str.len() != 11, "ERROR_SCHOOL_UDISE"] = 1

    # ---------- 2. DATE OF BIRTH Format/Missing Check ---------- #
    dob_str = df["DATE OF BIRTH"].astype(str)

    is_missing_dob = (
        df["DATE OF BIRTH"].isna()
        | dob_str.str.strip().eq("")
        | dob_str.str.lower().eq("nan")
    )

    contains_1_1 = dob_str.str.contains("1-1", case=False, na=False)
    contains_1jan = dob_str.str.contains("1jan", case=False, na=False)

    df.loc[is_missing_dob | contains_1_1 | contains_1jan, "ERROR_DOB_FORMAT"] = 1

    # ---------- 3. AGE Range Check (only ADOLOSCENT) ---------- #
    age_numeric = pd.to_numeric(df["AGE"], errors="coerce")
    program_subtype = df["PROGRAMSUBTYPENAME"].astype(str).str.upper().str.strip()

    is_adolescent = program_subtype.eq("ADOLOSCENT")
    age_out_range = (age_numeric <= 9) | (age_numeric >= 18)

    df.loc[is_adolescent & age_out_range, "ERROR_AGE_RANGE"] = 1

    # ---------- 4. Parent Consent Check ---------- #
    # ---------- 4. Parent Consent Check (NON-DESTRUCTIVE) ---------- #
    pc_raw = df["Parent Consent"]          # NEVER TOUCH THIS
    pc_work = pc_raw.astype(str).str.strip().str.lower()

    is_missing_pc = (
        pc_raw.isna() |
        pc_work.eq("") |
        pc_work.eq("nan")
    )

    is_no_pc = pc_work.eq("no")

    df.loc[is_missing_pc | is_no_pc, "ERROR_PARENT_CONSENT"] = 1

    # ---------- 5. P_Age Check (if ANY 0 exists → error) ---------- #
    p_age_str = df["P_Age"].astype(str)
    # \b0\b ensures 0 as a separate value (handles "0", "0, 35", "35, 0", "0, 0, 35", etc.)
    contains_zero = p_age_str.str.contains(r"\b0\b", regex=True, na=False)
    df.loc[contains_zero, "ERROR_P_AGE"] = 1

    # ---------- 6. RELIGIONNAME Missing Check ---------- #
    rel_str = df["RELIGIONNAME"].astype(str)

    is_missing_rel = (
        df["RELIGIONNAME"].isna()
        | rel_str.str.strip().eq("")
        | rel_str.str.lower().eq("nan")
        | rel_str.str.lower().eq("missing")
    )

    df.loc[is_missing_rel, "ERROR_RELIGIONNAME"] = 1

    # ---------- Total_Errors (sum of all flags) ---------- #
    error_cols = [
        "ERROR_SCHOOL_UDISE",
        "ERROR_DOB_FORMAT",
        "ERROR_AGE_RANGE",
        "ERROR_PARENT_CONSENT",
        "ERROR_P_AGE",
        "ERROR_RELIGIONNAME",
    ]
    df["Total_Errors"] = df[error_cols].sum(axis=1)

    return df


# --------------- Rules Sheet (Sheet Shee2) --------------- #
def build_rules_sheet() -> pd.DataFrame:
    """Create a DataFrame describing all validation rules."""
    data = [
        [
            "Rows removed",
            'Any row where any cell contains the text "Applied filters:" is removed before validation.',
        ],
        [
            "Dropped columns",
            "Columns DOCUMENTTYPE and DOCUMENTNO are dropped if present.",
        ],
        [
            "School UDISE",
            "ERROR_SCHOOL_UDISE = 1 when School UDISE (as text) length is not exactly 11 characters.",
        ],
        [
            "DATE OF BIRTH",
            "ERROR_DOB_FORMAT = 1 when DATE OF BIRTH is blank / NaN / 'nan' OR contains '1-1' OR contains '1Jan'.",
        ],
        [
            "AGE (ADOLOSCENT only)",
            'ERROR_AGE_RANGE = 1 when PROGRAMSUBTYPENAME = "ADOLOSCENT" and AGE <= 9 or AGE >= 18.',
        ],
        [
            "Parent Consent",
            "ERROR_PARENT_CONSENT = 1 when Parent Consent is blank / NaN / 'nan' OR equals 'No' (case-insensitive).",
        ],
        [
            "P_Age",
            "ERROR_P_AGE = 1 when P_Age contains any 0 value (e.g. '0', '0, 35', '35, 0', '0, 0, 40').",
        ],
        [
            "RELIGIONNAME",
            "ERROR_RELIGIONNAME = 1 when RELIGIONNAME is blank / NaN / 'nan' OR equals 'missing' (case-insensitive).",
        ],
        [
            "Total_Errors",
            "Total_Errors is the sum of all error flags: ERROR_SCHOOL_UDISE, ERROR_DOB_FORMAT, "
            "ERROR_AGE_RANGE, ERROR_PARENT_CONSENT, ERROR_P_AGE, ERROR_RELIGIONNAME.",
        ],
        [
            "Error_Tier (Gold/Silver/Bronze/Iron)",
            "Each row is assigned a quality tier based on its Total_Errors relative to the maximum Total_Errors "
            "in the file: if max Total_Errors = 0 → all Gold; else: 0 errors = Gold; "
            "0 < errors/max ≤ 0.33 = Silver; 0.33 < errors/max ≤ 0.66 = Bronze; > 0.66 = Iron.",
        ],
        [
            "ProgramLaunchName split (ZIP)",
            "The ZIP download output creates one Excel file per unique ProgramLaunchName "
            "from the validated data.",
        ],
    ]
    return pd.DataFrame(data, columns=["Check_Name", "Logic_Description"])


# --------------- Error Tier Classification --------------- #
def classify_error_tier(total_errors: int, max_errors: int) -> str:
    """Classify a row into Gold/Silver/Bronze/Iron based on Total_Errors."""
    if max_errors == 0:
        return "Gold"
    if total_errors == 0:
        return "Gold"
    ratio = total_errors / max_errors
    if ratio <= 0.33:
        return "Silver"
    elif ratio <= 0.66:
        return "Bronze"
    else:
        return "Iron"


# --------------- Helper: safe filename from ProgramLaunchName --------------- #
def safe_filename_from_pln(pln_value: str) -> str:
    text = str(pln_value).strip()
    # Replace bad characters
    safe = "".join(c if c.isalnum() or c in (" ", "_", "-") else "_" for c in text)
    safe = "_".join(safe.split())  # spaces -> single underscore
    return safe[:80] or "ProgramLaunchName"


# --------------- Streamlit App --------------- #
def main():
    st.set_page_config(
        page_title=f"CPRF Validation Tool | v{APP_VERSION}",
        page_icon="🔐",
        layout="wide",
    )

    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if not st.session_state["authenticated"]:
        render_login_page()
        return

    render_authenticated_header()

    page = st.radio(
        "Workspace",
        ["Validation", "User Manual", "Release Notes"],
        horizontal=True,
        label_visibility="collapsed",
    )
    st.divider()

    if page == "User Manual":
        render_manual_page()
        return
    if page == "Release Notes":
        render_release_notes_page()
        return

    # Main instructions + Total_Errors line
    st.write(
        """
Upload a CPRF `.xlsx` file with these mandatory columns:

- **School UDISE** (must be exactly 11 characters)
- **DATE OF BIRTH**
- **AGE**
- **PROGRAMSUBTYPENAME**
- **Parent Consent**
- **P_Age**
- **RELIGIONNAME**
- **ProgramLaunchName**

*(If your file has `DOCUMENTTYPE` or `DOCUMENTNO`, they will be dropped automatically.)*

**Validations performed:**

1. Removes rows containing `"Applied filters:"` in any column  
2. `ERROR_SCHOOL_UDISE = 1` → `School UDISE` length ≠ 11  
3. `ERROR_DOB_FORMAT = 1` → `DATE OF BIRTH` is blank / NaN / "nan" / contains `"1-1"` / `"1Jan"`  
4. `ERROR_AGE_RANGE = 1` → `PROGRAMSUBTYPENAME = "ADOLOSCENT"` and `AGE ≤ 9` or `AGE ≥ 18`  
5. `ERROR_PARENT_CONSENT = 1` → `Parent Consent` is blank / NaN / "nan" / "No"  
6. `ERROR_P_AGE = 1` → `P_Age` contains any `0` (even in comma-separated values)  
7. `ERROR_RELIGIONNAME = 1` → `RELIGIONNAME` is blank / NaN / "nan" / "missing"`  

**Total_Errors** = sum of all error flags for each row.
"""
    )

    # Score band table BEFORE explanation text
    st.markdown(
        """
<table style="width:70%; border-collapse: collapse; margin-left:auto; margin-right:auto; font-size:14px;">
    <tr style="border:1px solid #ddd; text-align:center; font-weight:bold;">
        <th style="border:1px solid #ddd; padding:8px;">Gold</th>
        <th style="border:1px solid #ddd; padding:8px;">Silver</th>
    </tr>
    <tr style="border:1px solid #ddd; text-align:left;">
        <td style="border:1px solid #ddd; padding:8px;">
            0 errors  
            <br>(or all rows if max errors = 0)
        </td>
        <td style="border:1px solid #ddd; padding:8px;">
            Low error count  
            <br>(0 &lt; errors/max ≤ 0.33)
        </td>
    </tr>
    <tr style="border:1px solid #ddd; text-align:center; font-weight:bold;">
        <th style="border:1px solid #ddd; padding:8px;">Bronze</th>
        <th style="border:1px solid #ddd; padding:8px;">Iron</th>
    </tr>
    <tr style="border:1px solid #ddd; text-align:left;">
        <td style="border:1px solid #ddd; padding:8px;">
            Medium error count  
            <br>(0.33 &lt; errors/max ≤ 0.66)
        </td>
        <td style="border:1px solid #ddd; padding:8px;">
            Highest error count  
            <br>(errors/max &gt; 0.66)
        </td>
    </tr>
</table>
""",
        unsafe_allow_html=True,
    )

    uploaded_file = st.file_uploader(
        "Upload CPRF Excel file (.xlsx)",
        type=["xlsx"],
        accept_multiple_files=False,
    )

    if uploaded_file is None:
        st.info("Please upload an Excel file to begin.")
        return

    st.success(f"File `{uploaded_file.name}` uploaded successfully.")

    try:
        with st.spinner("Reading and validating data..."):
            df = pd.read_excel(
                uploaded_file,
                dtype={"School UDISE": str}
            )
            processed_df = process_excel(df)

        # ---------- Add Error_Tier based on Total_Errors ----------
        max_errors = processed_df["Total_Errors"].max()
        processed_df["Error_Tier"] = processed_df["Total_Errors"].apply(
            lambda x: classify_error_tier(x, max_errors)
        )

        # ---------- Sort by Total_Errors (highest → lowest) ----------
        processed_df = processed_df.sort_values(
            by="Total_Errors", ascending=False
        ).reset_index(drop=True)

        st.success("Validation complete!")

        # ---------- SUMMARY METRICS ----------
        st.subheader("Summary")

        total_rows = len(processed_df)
        rows_with_errors = int((processed_df["Total_Errors"] > 0).sum())

        error_flag_cols = [
            "ERROR_SCHOOL_UDISE",
            "ERROR_DOB_FORMAT",
            "ERROR_AGE_RANGE",
            "ERROR_PARENT_CONSENT",
            "ERROR_P_AGE",
            "ERROR_RELIGIONNAME",
        ]
        flag_counts = {col: int(processed_df[col].sum()) for col in error_flag_cols}

        tier_counts = (
            processed_df["Error_Tier"]
            .value_counts()
            .reindex(["Gold", "Silver", "Bronze", "Iron"], fill_value=0)
        )

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("**Rows Summary**")
            st.write(f"- Total rows (after cleaning): **{total_rows}**")
            st.write(f"- Rows with `Total_Errors > 0`: **{rows_with_errors}**")

            st.markdown("**Rows by Quality Band (Error_Tier)**")
            for tier in ["Gold", "Silver", "Bronze", "Iron"]:
                st.write(f"- {tier}: **{tier_counts[tier]}** rows")

        with col2:
            st.markdown("**Error Counts by Category**")
            st.write(f"- ERROR_SCHOOL_UDISE: **{flag_counts['ERROR_SCHOOL_UDISE']}**")
            st.write(f"- ERROR_DOB_FORMAT: **{flag_counts['ERROR_DOB_FORMAT']}**")
            st.write(f"- ERROR_AGE_RANGE: **{flag_counts['ERROR_AGE_RANGE']}**")
            st.write(
                f"- ERROR_PARENT_CONSENT: **{flag_counts['ERROR_PARENT_CONSENT']}**"
            )
            st.write(f"- ERROR_P_AGE: **{flag_counts['ERROR_P_AGE']}**")
            st.write(
                f"- ERROR_RELIGIONNAME: **{flag_counts['ERROR_RELIGIONNAME']}**"
            )

        # --- PREVIEW (TOP 10 ERROR ROWS, already sorted by Total_Errors desc) ---
        st.subheader("Preview of Error Rows (Top 10 Only)")
        error_df = processed_df[processed_df["Total_Errors"] > 0].copy()

        if error_df.empty:
            st.info("No errors found! (Total_Errors = 0 for all rows)")
        else:
            st.dataframe(error_df.head(10))

        # --- MISSING CONTACT NUMBER FOLLOW-UP ---
        contact_columns = [
            "REGIONNAME",
            "STATENAME",
            "DISTRICTNAME",
            "Community/School",
            "School Type",
            "School UDISE",
            "PROGRAMTYPENAME",
            "PROGRAMSUBTYPENAME",
            "ProgramLaunchName",
            "FUNDERNAME",
            "Child School Name",
            "CHILDID",
            "CONTACTNUMBER",
        ]
        missing_contact_df = pd.DataFrame(columns=contact_columns + ["ContactNumberisMissing"])

        if "CONTACTNUMBER" in processed_df.columns:
            contact_work = processed_df["CONTACTNUMBER"].astype(str).str.strip()
            missing_contact_mask = (
                processed_df["CONTACTNUMBER"].isna()
                | contact_work.eq("")
                | contact_work.str.lower().isin(["nan", "none", "missing"])
            )
            available_contact_columns = [
                col for col in contact_columns if col in processed_df.columns
            ]
            missing_contact_df = processed_df.loc[
                missing_contact_mask, available_contact_columns
            ].copy()

            # Ensure the requested output columns always exist and remain in order.
            for col in contact_columns:
                if col not in missing_contact_df.columns:
                    missing_contact_df[col] = ""
            missing_contact_df = missing_contact_df[contact_columns]
            missing_contact_df["ContactNumberisMissing"] = "Phone Number Missing"

        st.subheader("Missing Contact Numbers")
        st.write(
            f"Records with missing CONTACTNUMBER: **{len(missing_contact_df):,}**"
        )
        if not missing_contact_df.empty:
            st.dataframe(missing_contact_df.head(20), use_container_width=True)
        else:
            st.info("No missing CONTACTNUMBER records found in this file.")

        # --- BUILD RULES SHEET DATAFRAME ---
        rules_df = build_rules_sheet()

        # --- DATE STRING FOR FILENAMES (v2026.01.02) ---
        # Output filenames use DDMMYYYY only; time and seconds are intentionally removed.
        base_name = Path(uploaded_file.name).stem
        file_date = datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%d%m%Y")

        validated_filename = f"{base_name}_Validated_{file_date}.xlsx"
        error_filename = f"{base_name}_Errors_{file_date}.xlsx"
        missing_contact_filename = (
            f"{base_name}_Missing_Contact_Numbers_{file_date}.xlsx"
        )
        zip_filename = f"{base_name}_ProgramLaunch_Files_{file_date}.zip"

        # --- DOWNLOAD: FULL DATASET ---
        full_output = io.BytesIO()
        with pd.ExcelWriter(full_output, engine="openpyxl") as writer:
            processed_df.to_excel(writer, index=False, sheet_name="Validated_Data")
            missing_contact_df.to_excel(
                writer, index=False, sheet_name="Contact_Number_Missing"
            )
            rules_df.to_excel(writer, index=False, sheet_name="Rules")
        full_output.seek(0)

        # --- DOWNLOAD: ERROR-ONLY DATASET ---
        error_output = io.BytesIO()
        with pd.ExcelWriter(error_output, engine="openpyxl") as writer:
            error_df.to_excel(writer, index=False, sheet_name="Error_Rows")
            rules_df.to_excel(writer, index=False, sheet_name="Rules")
        error_output.seek(0)

        # --- DOWNLOAD: MISSING CONTACT NUMBER DATASET ---
        missing_contact_output = io.BytesIO()
        with pd.ExcelWriter(missing_contact_output, engine="openpyxl") as writer:
            missing_contact_df.to_excel(
                writer, index=False, sheet_name="Missing_Contact_Numbers"
            )
        missing_contact_output.seek(0)

        # --- DOWNLOAD: ZIP BY ProgramLaunchName (full + error subset, per PLN) ---
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for pln, grp in processed_df.groupby("ProgramLaunchName"):
                # Full data for this ProgramLaunchName
                pln_full = grp.copy()
                # Error-only subset for this ProgramLaunchName
                pln_errors = grp[grp["Total_Errors"] > 0].copy()

                file_buffer = io.BytesIO()
                with pd.ExcelWriter(file_buffer, engine="openpyxl") as writer:
                    pln_full.to_excel(
                        writer,
                        index=False,
                        sheet_name="Validated_Data",
                    )
                    pln_errors.to_excel(
                        writer,
                        index=False,
                        sheet_name="Error_Rows",
                    )
                    rules_df.to_excel(writer, index=False, sheet_name="Rules")
                file_buffer.seek(0)

                safe_name = safe_filename_from_pln(pln)
                zipf.writestr(f"{safe_name}.xlsx", file_buffer.getvalue())

        zip_buffer.seek(0)

        st.subheader("Download Outputs")
        col_a, col_b, col_c, col_d = st.columns(4)

        with col_a:
            st.download_button(
                label="Download Full Validated Excel",
                data=full_output.getvalue(),
                file_name=validated_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        with col_b:
            st.download_button(
                label="Download Error Rows Only",
                data=error_output.getvalue(),
                file_name=error_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        with col_c:
            st.download_button(
                label="Download Missing Contact Numbers",
                data=missing_contact_output.getvalue(),
                file_name=missing_contact_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        with col_d:
            st.download_button(
                label="Download ZIP by ProgramLaunchName",
                data=zip_buffer.getvalue(),
                file_name=zip_filename,
                mime="application/x-zip-compressed",
            )

    except ValueError as ve:
        st.error(f"Error: {ve}")
    except Exception as e:
        st.error(f"Unexpected error: {e}")


if __name__ == "__main__":
    main()